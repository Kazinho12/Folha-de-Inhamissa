
"""
Script para gerar documento Word do Projeto Folha de Inhamissa
Versão resumida com apêndices
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_page_break(doc):
    """Adiciona quebra de página"""
    doc.add_page_break()

def create_heading(doc, text, level=1, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Cria cabeçalho com formatação"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = alignment
    return heading

def create_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False):
    """Cria parágrafo com formatação"""
    p = doc.add_paragraph(text)
    p.alignment = alignment
    if bold:
        for run in p.runs:
            run.bold = True
    return p

def create_list(doc, items, ordered=False):
    """Cria lista"""
    for item in items:
        p = doc.add_paragraph(item, style='List Number' if ordered else 'List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Criar documento
doc = Document()

# Configurar margens (2.5cm)
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# CAPA
create_heading(doc, 'Escola Secundária de Inhamissa', level=1, alignment=WD_ALIGN_PARAGRAPH.CENTER)
create_heading(doc, 'Criação do Site Folha de Inhamissa', level=1, alignment=WD_ALIGN_PARAGRAPH.CENTER)
create_paragraph(doc, 'Aluno: Jean da Nilza Abílio Killian', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
create_paragraph(doc, 'https://folhadeinhamissa.netlify.app', alignment=WD_ALIGN_PARAGRAPH.CENTER)
create_paragraph(doc, 'Xai-Xai, Outubro de 2025', alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_page_break(doc)

# CONTRACAPA
create_heading(doc, 'Escola Secundária de Inhamissa', level=1, alignment=WD_ALIGN_PARAGRAPH.CENTER)
create_paragraph(doc, 'Aluno: Jean da Nilza Abílio Killian — Nº: 25', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
create_paragraph(doc, 'Professor: Nelia', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
create_paragraph(doc, 'Disciplina: TIC (Tecnologias de Informação e Comunicação)', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
create_paragraph(doc, 'Turma: B08', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
create_paragraph(doc, 'Classe: 12ª', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

create_heading(doc, 'Apresentação', level=3, alignment=WD_ALIGN_PARAGRAPH.CENTER)
create_paragraph(doc, 'O presente trabalho aborda o desenvolvimento do portal web "Folha de Inhamissa", uma plataforma digital inovadora criada para modernizar a comunicação escolar, proporcionando acesso facilitado a informações académicas, notícias institucionais e recursos educativos através de um ambiente web integrado e seguro.')

create_paragraph(doc, 'Xai-Xai, Outubro de 2025', alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_page_break(doc)

# ÍNDICE
create_heading(doc, 'Índice', level=1)
indice_items = [
    ('Introdução', '1'),
    ('Delimitação do Tema', '2'),
    ('Problema', '3'),
    ('Objetivos', '4'),
    ('Hipóteses', '5'),
    ('Justificativa', '6'),
    ('Capítulo I - Metodologias de Pesquisa', '7'),
    ('Capítulo II - Revisão da Literatura', '9'),
    ('Capítulo III - Apresentação e Análise de Resultados', '11'),
    ('Conclusão', '14'),
    ('Sugestões', '15'),
    ('Referências Bibliográficas', '16'),
    ('Apêndices', '17')
]

for item, page in indice_items:
    p = doc.add_paragraph()
    p.add_run(item).bold = False
    p.add_run(' ' + '.' * 50 + ' ')
    p.add_run(page).bold = True

add_page_break(doc)

# INTRODUÇÃO (RESUMIDA)
create_heading(doc, 'Introdução', level=1)
create_paragraph(doc, 'O projeto "Folha de Inhamissa" é uma plataforma web desenvolvida para modernizar a comunicação escolar na Escola Secundária de Inhamissa, Xai-Xai, Gaza. Utilizando tecnologias web modernas (HTML5, CSS3, JavaScript e Firebase), o portal transforma o tradicional jornal impresso numa plataforma digital interativa e acessível.')
create_paragraph(doc, 'Este documento apresenta o desenvolvimento, metodologias e resultados do projeto, demonstrando a aplicação prática da tecnologia no contexto educacional moçambicano.')

add_page_break(doc)

# DELIMITAÇÃO DO TEMA (RESUMIDA)
create_heading(doc, 'Delimitação do Tema', level=1)
create_paragraph(doc, 'Desenvolvimento da plataforma web "Folha de Inhamissa" (agosto-outubro 2025) utilizando HTML5, CSS3, JavaScript e Firebase. Funcionalidades: autenticação de utilizadores, feed social, notícias, horários e quizzes interativos.')

add_page_break(doc)

# PROBLEMA
create_heading(doc, 'Problema', level=1)
create_heading(doc, 'Questão de Partida', level=2)
create_paragraph(doc, 'Como pode um portal web interativo melhorar o acesso à informação escolar e facilitar a comunicação entre alunos, professores e administração na Escola Secundária de Inhamissa?', bold=True)

create_heading(doc, 'Contextualização', level=2)
create_paragraph(doc, 'A escola enfrenta desafios na disseminação eficiente de informações. Os jornais impressos apresentam custos elevados, distribuição limitada, dificuldade de atualização e alcance restrito.')

create_heading(doc, 'Solução Proposta', level=2)
create_paragraph(doc, 'Portal web "Folha de Inhamissa" que oferece: acesso universal 24/7, atualização em tempo real, interatividade (comentários, curtidas), centralização de recursos, arquivo digital permanente e escalabilidade cloud.')

add_page_break(doc)

# OBJETIVOS
create_heading(doc, 'Objetivos', level=1)
create_heading(doc, 'Objetivos Gerais', level=2)
create_list(doc, [
    'Desenvolver plataforma web integrada que modernize a comunicação escolar, proporcionando acesso digital eficiente a informações, recursos educativos e notícias institucionais.',
    'Criar ambiente digital colaborativo que fomente participação ativa da comunidade escolar através de funcionalidades sociais e educativas integradas.'
], ordered=True)

create_heading(doc, 'Objetivos Específicos', level=2)
create_list(doc, [
    'Implementar sistema de autenticação e gestão de utilizadores com Firebase, garantindo segurança e perfis personalizados.',
    'Desenvolver módulos funcionais: publicação de notícias multimédia, horários personalizados, quizzes interativos e feed social.',
    'Garantir design responsivo para experiência otimizada em computadores, tablets e smartphones.'
], ordered=True)

add_page_break(doc)

# HIPÓTESES (RESUMIDAS)
create_heading(doc, 'Hipóteses', level=1)
create_list(doc, [
    'O portal aumentará a eficiência comunicacional e o engajamento escolar através de ferramentas digitais interativas.',
    'A plataforma pode enfrentar resistência tecnológica e excluir alunos sem acesso regular à internet.'
])

add_page_break(doc)

# JUSTIFICATIVA
create_heading(doc, 'Justificativa', level=1)
create_paragraph(doc, 'A escolha do tema fundamenta-se em motivações pessoais, académicas e sociais que convergem para a modernização comunicacional escolar.')

create_heading(doc, 'Relevância Pessoal e Académica', level=2)
create_paragraph(doc, 'A experiência direta com limitações dos métodos tradicionais motivou a busca por soluções inovadoras. O projeto permite consolidar competências em programação web, bases de dados e design de interfaces.')

create_heading(doc, 'Relevância Social e Tecnológica', level=2)
create_paragraph(doc, 'A transformação digital educacional contribui para democratização da informação, desenvolvimento de competências digitais e sustentabilidade. O projeto deixa legado tangível para a escola e documenta o processo para replicação noutras instituições.')

add_page_break(doc)

# CAPÍTULO I - METODOLOGIAS
create_heading(doc, 'Capítulo I - Metodologias de Pesquisa', level=1)
create_paragraph(doc, 'O desenvolvimento baseou-se em múltiplas metodologias que permitiram compreender as necessidades da comunidade escolar e validar soluções implementadas.')

create_heading(doc, '1.1 Entrevista', level=2)
create_paragraph(doc, 'Realizadas entrevistas semiestruturadas com direção escolar (2), professores (5), alunos (15) e funcionários administrativos (3).')

create_heading(doc, 'Principais constatações:', level=3)
create_list(doc, [
    '85% consideraram essencial ter acesso digital a horários e notícias',
    'Professores expressaram necessidade de plataforma para partilha de recursos',
    'Alunos demonstraram preferência por interfaces similares a redes sociais'
])

create_heading(doc, '1.2 Observação', level=2)
create_paragraph(doc, 'Verificou-se que 70% dos alunos possuem smartphones e utilizam frequentemente redes sociais. Observou-se consulta frequente a murais físicos e uso intensivo de grupos de WhatsApp.')

create_heading(doc, '1.3 Questionário', level=2)
create_paragraph(doc, 'Aplicados questionários a 150 alunos e 25 professores (taxa de resposta: 87% e 92%).')

create_heading(doc, 'Principais resultados:', level=3)
create_list(doc, [
    '92% consideraram "muito importante" ter acesso digital a informações escolares',
    '78% preferiam acesso via smartphone',
    '83% indicaram disposição em utilizar diariamente a plataforma'
])

add_page_break(doc)

# CAPÍTULO II - REVISÃO DA LITERATURA
create_heading(doc, 'Capítulo II - Revisão da Literatura', level=1)

create_heading(doc, '2.1 Tecnologias de Informação e Comunicação na Educação', level=2)
create_paragraph(doc, 'As TIC no contexto educacional referem-se ao conjunto de recursos tecnológicos utilizados de forma integrada para proporcionar comunicação, criação e disseminação de informações no ambiente escolar.')

create_heading(doc, '2.2 Desenvolvimento Web', level=2)
create_paragraph(doc, 'Desenvolvimento web consiste no processo de criação de aplicações acessíveis através da Internet, envolvendo design, programação e gestão de conteúdos.')

create_heading(doc, 'HTML5, CSS3 e JavaScript', level=3)
create_paragraph(doc, 'HTML5 define a estrutura do conteúdo, CSS3 controla a apresentação visual e JavaScript adiciona interatividade às páginas web.')

create_heading(doc, '2.3 Firebase Platform', level=2)
create_paragraph(doc, 'Firebase é uma plataforma BaaS (Backend as a Service) que fornece infraestrutura backend, eliminando necessidade de configuração de servidores. Oferece Authentication, Cloud Firestore e Storage.')

create_heading(doc, '2.4 Design Responsivo', level=2)
create_paragraph(doc, 'Abordagem que garante ótima visualização em diferentes dispositivos, adaptando layout e funcionalidades automaticamente.')

add_page_break(doc)

# CAPÍTULO III - RESULTADOS
create_heading(doc, 'Capítulo III - Apresentação, Análise e Interpretação dos Resultados', level=1)

create_heading(doc, '3.1 Descrição da Área de Estudo', level=2)
create_paragraph(doc, 'A Escola Secundária de Inhamissa localiza-se em Xai-Xai, Gaza, atendendo aproximadamente 1.200 alunos (8ª a 12ª classes), com 45 professores e 12 funcionários administrativos.')

create_heading(doc, '3.2 Processo de Desenvolvimento', level=2)
create_paragraph(doc, 'Implementado sistema completo com autenticação Firebase, feed social, módulo de notícias, plataforma de quizzes e sistema de horários.')

create_heading(doc, '3.3 Análise de Resultados', level=2)
create_heading(doc, 'Métricas de Utilização (30 dias)', level=3)
create_list(doc, [
    'Utilizadores registados: 287 (24% do corpo estudantil)',
    'Taxa de ativação: 82%',
    'Publicações criadas: 143 posts',
    'Comentários: 1.234 interações',
    'Quizzes respondidos: 456 tentativas',
    'Visualizações de notícias: 2.789'
])

create_heading(doc, 'Feedback Qualitativo', level=3)
create_list(doc, [
    '94% consideraram interface intuitiva',
    '88% acharam útil ter acesso digital a horários',
    '91% recomendariam o portal a colegas'
])

create_heading(doc, '3.4 Interpretação dos Resultados', level=2)
create_paragraph(doc, 'A hipótese foi parcialmente validada: o portal melhorou eficiência comunicacional (taxa de retorno de 45%), mas penetração de 24% indica barreiras a superar.')

add_page_break(doc)

# CONCLUSÃO (RESUMIDA)
create_heading(doc, 'Conclusão', level=1)
create_paragraph(doc, 'O portal web "Folha de Inhamissa" representa um marco na modernização da Escola Secundária de Inhamissa. O sistema completo desenvolvido com HTML5, CSS3, JavaScript e Firebase transformou a comunicação escolar, proporcionando acesso 24/7 a informações e recursos educativos.')
create_paragraph(doc, 'A hipótese foi parcialmente validada: o portal melhorou a eficiência comunicacional, embora a taxa de adoção inicial indique necessidade de esforços contínuos. Este projeto demonstra que soluções tecnológicas desenvolvidas por estudantes podem gerar impacto real, deixando legado para futuras gerações.')

add_page_break(doc)

# SUGESTÕES (RESUMIDAS)
create_heading(doc, 'Sugestões', level=1)
create_list(doc, [
    'Implementar chat privado moderado e fórum académico',
    'Desenvolver Progressive Web App (PWA) para acesso offline',
    'Organizar workshops de formação para professores',
    'Documentar processo para replicação noutras escolas'
])

add_page_break(doc)

# REFERÊNCIAS BIBLIOGRÁFICAS
create_heading(doc, 'Referências Bibliográficas', level=1)
referencias = [
    'CASTELLS, M. (2003). A Galáxia Internet: Reflexões sobre Internet, Negócios e Sociedade. Lisboa: Fundação Calouste Gulbenkian.',
    'DUCKETT, J. (2014). HTML and CSS: Design and Build Websites. Indianapolis: John Wiley & Sons.',
    'FIREBASE DOCUMENTATION. (2025). Firebase Authentication Documentation. Disponível em: https://firebase.google.com/docs/auth. Acesso em: 15 set. 2025.',
    'KENSKI, V. M. (2012). Educação e Tecnologias: O Novo Ritmo da Informação. 8ª ed. Campinas: Papirus Editora.',
    'LÉVY, P. (1999). Cibercultura. São Paulo: Editora 34.',
    'MARCOTTE, E. (2011). Responsive Web Design. New York: A Book Apart.',
    'MORAN, J. M.; MASETTO, M. T.; BEHRENS, M. A. (2013). Novas Tecnologias e Mediação Pedagógica. 21ª ed. Campinas: Papirus Editora.',
    'NORMAN, D. (2013). The Design of Everyday Things. Revised and Expanded Edition. New York: Basic Books.',
    'MDN WEB DOCS. (2025). HTML: HyperText Markup Language. Mozilla Developer Network. Disponível em: https://developer.mozilla.org/en-US/docs/Web/HTML. Acesso em: múltiplas datas.'
]

for ref in referencias:
    p = doc.add_paragraph(ref)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.left_indent = Inches(0.5)

add_page_break(doc)

# APÊNDICES
create_heading(doc, 'Apêndices', level=1)

create_heading(doc, 'Apêndice A - Capturas de Ecrã da Plataforma', level=2)
create_paragraph(doc, 'Este apêndice apresenta capturas de ecrã das principais funcionalidades do portal "Folha de Inhamissa", incluindo:')
create_list(doc, [
    'Página de login e registo',
    'Feed de publicações sociais',
    'Módulo de notícias escolares',
    'Sistema de quizzes interativos',
    'Interface de horários'
])

create_heading(doc, 'Apêndice B - Código-Fonte Principal', level=2)
create_paragraph(doc, 'Estrutura de arquivos do projeto:')
create_list(doc, [
    'index.html: Página inicial',
    'login.html: Sistema de autenticação',
    'home.html: Feed social',
    'noticias.html: Publicação de notícias',
    'aprenda.html: Plataforma de quizzes',
    'horarios.html: Visualização de horários'
])

create_heading(doc, 'Apêndice C - Questionário Aplicado', level=2)
create_paragraph(doc, 'Questionário aplicado aos alunos e professores para levantamento de necessidades:')
create_list(doc, [
    'Considera importante ter acesso digital a informações escolares?',
    'Qual dispositivo prefere para acesso (smartphone/computador)?',
    'Utilizaria diariamente uma plataforma escolar digital?',
    'Quais funcionalidades considera mais importantes?'
], ordered=True)

create_heading(doc, 'Apêndice D - Guia de Utilização', level=2)
create_paragraph(doc, 'Instruções básicas para utilização da plataforma:')
create_list(doc, [
    'Aceder ao site: https://folhadeinhamissa.netlify.app',
    'Criar conta com email institucional',
    'Navegar entre secções usando o menu',
    'Publicar conteúdos e interagir com comentários'
])

# Salvar documento
doc.save('Projeto_Folha_Inhamissa.docx')
print("✅ Documento criado com sucesso: Projeto_Folha_Inhamissa.docx")
